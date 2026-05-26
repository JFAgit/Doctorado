from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(r"C:\Users\fran_\Documents\Doctorado\Inicios\Dani\GSE30595_M1_2026\06_datasets_externos\GMCSF_basales_busqueda")
DOCX = OUT_DIR / "reporte_GMCSF_MDM_basales_para_director.docx"
TXT = OUT_DIR / "reporte_GMCSF_MDM_basales_para_director.txt"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EF", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, grid=None):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    if grid:
        old_grid = tbl.tblGrid
        if old_grid is not None:
            tbl.remove(old_grid)
        tbl_grid = OxmlElement("w:tblGrid")
        for w in grid:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(w))
            tbl_grid.append(col)
        tbl.insert(0, tbl_grid)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(grid[idx]))


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 58, 95)
    title.paragraph_format.space_after = Pt(4)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_key_value(doc, label, value):
    p = doc.add_paragraph()
    r = p.add_run(label + ": ")
    r.bold = True
    p.add_run(value)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_table_width(table, grid=widths)

    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], "F2F4F7")
        set_cell_margins(hdr[i])
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
    set_table_width(table, grid=widths)
    return table


def build_docx():
    doc = Document()
    style_doc(doc)

    doc.add_paragraph("Resumen de datasets GM-CSF MDM basales humanos", style="Title")
    subtitle = doc.add_paragraph()
    subtitle.add_run("Busqueda para analisis por sexo de expresion inmunometabolica, proinflamatoria y antiinflamatoria/M2.").italic = True

    add_key_value(doc, "Criterio principal", "Macrofagos humanos derivados de monocitos de sangre periferica, diferenciados con GM-CSF, en condicion basal/unstimulated/untreated/control.")
    add_key_value(doc, "Prioridad", "Maximizar n femenino y masculino sin contar tiempos repetidos como individuos independientes.")

    doc.add_heading("Conclusion principal", level=1)
    p = doc.add_paragraph()
    p.add_run("Recomendacion para analisis principal: ").bold = True
    p.add_run("usar 18 muestras/donantes GM-CSF MDM basales con balance final 9F / 9M.")
    p = doc.add_paragraph()
    p.add_run("Datasets incluidos: ").bold = True
    p.add_run("GSE160862 + GSE160863 + GSE232044 + GSE224845, contando un solo basal UNT por donante en GSE224845.")

    doc.add_heading("Conteo final por estrategia", level=1)
    rows = [
        ["Nucleo mas limpio", "GSE160862 + GSE160863", 6, 4, 10, "Basal/unactivated, GM-CSF MDM, healthy donors."],
        ["Nucleo + balance", "GSE160862 + GSE160863 + GSE232044", 8, 6, 14, "Agrega untreated GM-CSF MDM; queda bastante balanceado."],
        ["Recomendado estricto", "GSE160862 + GSE160863 + GSE232044 + GSE224845, usando 1 UNT por donante", 9, 9, 18, "Mejor compromiso: GM-CSF basal y sexos balanceados."],
        ["Sensibilidad", "Recomendado estricto + GSE266236 siRNA control", 10, 11, 21, "Suma controles siRNA; util, pero mas heterogeneo."],
        ["Todo lo usable con sexo inferido", "Recomendado + GSE266236 + GSE304218 + GSE102492 + GSE156696 + GSE256208", 10, 30, 40, "Mayor n, pero muy masculino y mezcla vehiculos/controles tecnicos."],
    ]
    add_table(
        doc,
        ["Estrategia", "Datasets incluidos", "F", "M", "Total", "Comentario"],
        rows,
        [1550, 3350, 520, 520, 650, 2770],
    )

    doc.add_heading("Datasets estrictos mas relevantes", level=1)
    rows2 = [
        ["GSE160862 + GSE160863", "10 donantes basales", "6F / 4M", "Mejor candidato estricto."],
        ["GSE232044", "4 untreated", "2F / 2M", "Muy limpio y balanceado, n chico."],
        ["GSE224845", "4 donantes UNT; 12 muestras si se modela tiempo", "1F / 3M", "Usar un solo tiempo por donante para n independiente."],
        ["GSE266236", "3 siRNA control", "1F / 2M", "Opcion de sensibilidad."],
        ["GSE304218", "5 controles no infectados", "0F / 5M", "Basal usable, pero solo masculino."],
        ["GSE102492", "8 GM-CSF MDM", "0F / 8M", "N bueno, pero solo masculino."],
    ]
    add_table(doc, ["Dataset", "Basales GM-CSF usables", "Sexo", "Uso sugerido"], rows2, [2100, 2500, 1200, 3560])

    doc.add_heading("Nota sobre el dataset grande", level=1)
    p = doc.add_paragraph()
    p.add_run("GSE269009 ").bold = True
    p.add_run("tiene n alto y sexo explicito, pero no cumple el criterio GM-CSF MDM estricto: los macrofagos MP fueron diferenciados con M-CSF; GM-CSF + IL4 se uso para dendriticas. Si se flexibiliza el criterio a M-CSF-derived macrophages, aporta 109 MP: controles 34F / 23M y casos 25F / 27M.")

    doc.add_heading("Decision sugerida", level=1)
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Analisis principal: ").bold = True
    p.add_run("18 muestras/donantes, 9F / 9M.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Analisis de sensibilidad: ").bold = True
    p.add_run("21 muestras, 10F / 11M, agregando GSE266236.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("No recomendado como n independiente: ").bold = True
    p.add_run("contar los 12 tiempos UNT de GSE224845 como 12 individuos; si se usan todos, modelar donor/tiempo.")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("GM-CSF MDM basales - resumen interno").font.size = Pt(8)

    doc.save(DOCX)


def build_txt():
    text = """Resumen de datasets GM-CSF MDM basales humanos

Criterio principal:
Macrofagos humanos derivados de monocitos de sangre periferica, diferenciados con GM-CSF, en condicion basal/unstimulated/untreated/control.

Conclusion principal:
Para el analisis principal se recomienda usar 18 muestras/donantes GM-CSF MDM basales con balance final 9F / 9M.

Datasets incluidos en el recomendado estricto:
- GSE160862 + GSE160863
- GSE232044
- GSE224845, contando un solo basal UNT por donante

Conteo final por estrategia:
1. Nucleo mas limpio: GSE160862 + GSE160863
   F = 6, M = 4, total = 10
   Basal/unactivated, GM-CSF MDM, healthy donors.

2. Nucleo + balance: GSE160862 + GSE160863 + GSE232044
   F = 8, M = 6, total = 14
   Agrega untreated GM-CSF MDM; queda bastante balanceado.

3. Recomendado estricto: GSE160862 + GSE160863 + GSE232044 + GSE224845, usando 1 UNT por donante
   F = 9, M = 9, total = 18
   Mejor compromiso: GM-CSF basal y sexos balanceados.

4. Sensibilidad: recomendado estricto + GSE266236 siRNA control
   F = 10, M = 11, total = 21
   Suma controles siRNA; util, pero mas heterogeneo.

5. Todo lo usable con sexo inferido:
   Recomendado + GSE266236 + GSE304218 + GSE102492 + GSE156696 + GSE256208
   F = 10, M = 30, total = 40
   Mayor n, pero muy masculino y mezcla vehiculos/controles tecnicos.

Datasets estrictos mas relevantes:
- GSE160862 + GSE160863: 10 donantes basales, 6F / 4M. Mejor candidato estricto.
- GSE232044: 4 untreated, 2F / 2M. Muy limpio y balanceado, n chico.
- GSE224845: 4 donantes UNT, 1F / 3M. Usar un solo tiempo por donante para n independiente.
- GSE266236: 3 siRNA control, 1F / 2M. Opcion de sensibilidad.
- GSE304218: 5 controles no infectados, 0F / 5M. Basal usable, pero solo masculino.
- GSE102492: 8 GM-CSF MDM, 0F / 8M. N bueno, pero solo masculino.

Nota sobre GSE269009:
GSE269009 tiene n alto y sexo explicito, pero no cumple el criterio GM-CSF MDM estricto: los macrofagos MP fueron diferenciados con M-CSF; GM-CSF + IL4 se uso para dendriticas. Si se flexibiliza el criterio a M-CSF-derived macrophages, aporta 109 MP: controles 34F / 23M y casos 25F / 27M.

Decision sugerida:
- Analisis principal: 18 muestras/donantes, 9F / 9M.
- Analisis de sensibilidad: 21 muestras, 10F / 11M, agregando GSE266236.
- No contar los 12 tiempos UNT de GSE224845 como 12 individuos; si se usan todos, modelar donor/tiempo.
"""
    TXT.write_text(text, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_txt()
    print(DOCX)
    print(TXT)
